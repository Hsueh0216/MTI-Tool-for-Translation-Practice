import streamlit as st
import fitz  # PyMuPDF
from docx import Document
from openai import OpenAI
from google import genai
from deep_translator import GoogleTranslator
import pandas as pd
import re
import time
import json
import io

# ================= 页面全局设置 =================
st.set_page_config(page_title="MTI 翻译实践小助手", page_icon="🎓", layout="wide")

# ================= 核心状态机初始化 =================
# 用于实现“断点续传”和“过程文件留存”
if 'doc_states' not in st.session_state:
    st.session_state.doc_states = {}

# ================= 核心工具函数 =================
def clean_xml_chars(text):
    if not isinstance(text, str): return str(text)
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

def call_llm(provider, api_key, system_prompt, user_prompt, temperature=0.1):
    """底层大模型统一路由 (已将超时放宽至 150 秒)"""
    if provider == "DeepSeek":
        client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com", timeout=150.0)
        res = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=temperature
        )
        return res.choices[0].message.content.strip()
    elif provider == "OpenAI":
        client = OpenAI(api_key=api_key, timeout=150.0)
        res = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=temperature
        )
        return res.choices[0].message.content.strip()
    elif provider == "Gemini":
        client = genai.Client(api_key=api_key)
        full_prompt = f"System Instruction: {system_prompt}\n\nUser Input: {user_prompt}"
        res = client.models.generate_content(model='gemini-1.5-flash', contents=full_prompt)
        return res.text.strip()
    return ""

def parse_termbase(file_stream) -> dict:
    try:
        df = pd.read_excel(file_stream)
        df.columns = df.columns.str.strip()
        if "Source" not in df.columns or "Target" not in df.columns: return {}
        df = df.dropna(subset=['Source', 'Target'])
        return dict(zip(df['Source'].astype(str).str.strip(), df['Target'].astype(str).str.strip()))
    except: return {}

# 【优化版】：自动抽取术语库（强规则过滤 & 扩大提取量）
def extract_auto_terms(paragraphs: list, target_lang: str, provider: str, api_key: str) -> dict:
    # 1. 扩大采样范围：将前 4000 字符提升到 10000 字符，跳过通常缺乏术语的“摘要/作者简介”部分
    sample_text = "\n".join(paragraphs)[:10000] 
    
    # 2. 强力 Prompt：明确要求提取数量，并使用【负面清单】严格禁止提取人名和书名
    sys_prompt = f"""你是一位极其严谨的学术译员和术语管理专家（Terminologist）。
    请从以下文本中提取 30 到 50 个最具代表性的【核心专业术语】。

    【核心筛选规则（极其重要）】：
    1. 必须是特定学科的理论概念、专业名词、核心方法论或行业黑话（Jargon）。
    2. 🚫 绝对禁止提取：人名（如学者名/作者名）、书名、文章标题、期刊名、出版地、机构名称、年份。
    3. 🚫 绝对禁止提取：日常通用词汇（如 research, study, analysis 等无门槛词汇）。
    4. 请将其精准、符合学术规范地翻译为{target_lang}。

    请严格输出合法的 JSON 数组格式，绝对不要包含任何其他多余的解释文字，格式如下：
    [
        {{"Source": "英文专业术语1", "Target": "中文专业译名1"}},
        {{"Source": "英文专业术语2", "Target": "中文专业译名2"}}
    ]"""
    
    for attempt in range(3):
        try:
            # temperature 设为 0.1，让大模型在遵守规则和输出 JSON 时保持最高级别的理性和严谨
            res = call_llm(provider, api_key, sys_prompt, sample_text, temperature=0.1)
            json_match = re.search(r'\[.*\]', res, re.DOTALL)
            if json_match:
                term_list = json.loads(json_match.group(0))
                # 二次清洗：剔除空值或明显不合理的短数据
                filtered_terms = {
                    item["Source"].strip(): item["Target"].strip() 
                    for item in term_list 
                    if "Source" in item and "Target" in item 
                    and len(item["Source"].strip()) > 1
                }
                return filtered_terms
        except Exception as e:
            if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e): 
                time.sleep(15)
            else: 
                continue
    return {}

def dict_to_excel(term_dict: dict) -> io.BytesIO:
    df = pd.DataFrame(list(term_dict.items()), columns=["Source", "Target"])
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
    output.seek(0)
    return output

def paragraphs_to_word(paragraphs: list) -> io.BytesIO:
    doc = Document()
    doc.add_heading('阶段一：清洗后原文提取', 0)
    for p in paragraphs: doc.add_paragraph(p)
    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def generate_mti_report(bilingual_pairs: list, termbase_dict: dict, theory: str, provider: str, api_key: str, status_placeholder=None) -> str:
    """【阶段三进阶版】采用分步组装策略 (Map-Reduce)，强迫大模型进行多轮扩写，生成深度过万字的 MTI 实践报告"""
    
    # 1. 组装语料样本 (依然保留截断防御)
    sample_texts = ""
    char_count = 0
    for pair in bilingual_pairs:
        chunk = f"【原文】{pair['source']}\n【译文】{pair['target']}\n\n"
        sample_texts += chunk
        char_count += len(chunk)
        if char_count > 8000: break

    term_str = "\n".join([f"{k} -> {v}" for k, v in termbase_dict.items()]) if termbase_dict else "无术语提取"

    # 2. 建立四段式独立 Prompt 矩阵 (逼迫模型对每一节进行深度发散)
    prompts = [
        (
            "一、 翻译项目概述与文本特征分析",
            f"请基于以下双语语料样本，撰写翻译实践报告的【第一部分】。\n要求：详尽分析源文本的语言风格、专业领域、词法（如专有名词、长难句）与句法特点，以及由此带来的总体翻译难点。字数要求 800-1000 字。严禁输出其他章节的内容。\n\n语料样本：\n{sample_texts}"
        ),
        (
            "二、 术语管理与验证",
            f"请撰写翻译实践报告的【第二部分】。\n请基于以下核心术语表，详细评估本次翻译中术语库的执行情况。请至少选取 4 个核心术语，深度剖析其翻译策略（如直译、意译、增词、转换等）及其对提升文本专业性的贡献。字数要求 800-1000 字。\n\n术语表：\n{term_str}"
        ),
        (
            f"三、 基于【{theory}】的案例分析",
            f"这是本报告的最核心章节。请基于以下双语语料样本和【{theory}】的理论框架，撰写报告的【第三部分】。\n要求：精准抽取 4-5 个最具代表性的长难句或特殊表达案例。每个案例必须独立成段并包含：\n1. 原译文对照\n2. 翻译难点深度剖析\n3. 严谨的学理分析（明确指出具体的翻译技巧，并用【{theory}】的核心概念论证“为何如此翻译”）。\n本部分字数要求不少于 1500 字，必须极具学术深度。\n\n语料样本：\n{sample_texts}"
        ),
        (
            "四、 翻译项目复盘与反思",
            f"请结合上述关于【{theory}】的翻译实践，撰写报告的【第四部分】。\n要求：深刻总结机器翻译（MT）在此类文本中的局限性、本地化术语库强干预的实际效果，以及作为译后编辑（MTPE）在双语能力和理论运用层面的收获。字数要求 600-800 字。"
        )
    ]

    full_report_md = ""
    base_system_prompt = "你是一位拥有深厚学术背景的 MTI（翻译硕士）导师及资深学术期刊审稿人。请严格使用学术书面语，逻辑严密，杜绝任何 AI 常见的口语化或套话表达。"

    # 3. 循环遍历，依次生成每一章并拼接
    for idx, (section_title, user_prompt) in enumerate(prompts):
        # 动态更新前端界面的状态文字
        if status_placeholder:
            status_placeholder.update(label=f"【阶段三】正在深度撰写：{section_title} ({idx+1}/4)...", state="running")
            
        for attempt in range(3):
            try:
                # 稍微调高 temperature (0.5)，让模型在长篇论述时更有发散性和文采
                section_content = call_llm(provider, api_key, base_system_prompt, user_prompt, temperature=0.5)
                
                # 清洗大模型可能自带的 markdown 代码块外壳
                section_content = re.sub(r'^```markdown|```$', '', section_content.strip(), flags=re.MULTILINE)
                
                # 拼接入总报告
                full_report_md += f"## {section_title}\n\n{section_content}\n\n---\n\n"
                
                time.sleep(2) # 章节之间的缓冲防并发限流
                break
            except Exception as e:
                if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e):
                    time.sleep(20)
                else:
                    raise e

    return full_report_md

def _add_formatted_runs(paragraph, text):
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 != 0: run.bold = True

def markdown_to_word(md_text: str, theory: str) -> io.BytesIO:
    doc = Document()
    md_text = re.sub(r'```markdown|```', '', md_text).strip()
    title = doc.add_heading(f'翻译实践报告：基于{theory}', 0)
    title.alignment = 1 
    for line in md_text.split('\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('### '): doc.add_heading(line[4:], level=3)
        elif line.startswith('## '): doc.add_heading(line[3:], level=2)
        elif line.startswith('# '): doc.add_heading(line[2:], level=1)
        elif line.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_runs(p, line[2:])
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            _add_formatted_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_formatted_runs(p, line)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


# ================= UI 界面绘制 =================
st.title("🎓 MTI 翻译实践小助手 (Pro版)")

with st.sidebar:
    st.header("⚙️ 引擎设置")
    ai_provider = st.selectbox("核心引擎", ["DeepSeek", "OpenAI", "Gemini"])
    api_key = st.text_input(f"请输入 {ai_provider} API Key", type="password")
    target_lang = st.selectbox("目标语言", ["简体中文", "English", "日本語"])
    
    st.divider()
    st.header("🛠️ 进阶功能")
    auto_term = st.checkbox("🤖 智能抽取术语库 (翻译前执行)", value=True, help="大模型自动提取专有名词并生成 Excel 术语库")
    enable_report = st.checkbox("📝 自动生成实践报告", value=True)
    translation_theory = st.selectbox("案例分析理论", ["目的论 (Skopos Theory)", "交际翻译与语义翻译 (Newmark)", "功能对等理论 (Nida)", "文本类型理论 (Reiss)"])

col1, col2 = st.columns(2)
with col1:
    termbase_file = st.file_uploader("导入已有术语库 (.xlsx, 可选)", type=['xlsx'])
    user_termbase = parse_termbase(termbase_file) if termbase_file else {}
with col2:
    uploaded_files = st.file_uploader("待翻译文档 (支持多文件与断点续传)", type=['pdf', 'docx'], accept_multiple_files=True)

# ================= 核心处理流 (带断点续传状态机) =================
if st.button("🚀 开始 / 继续处理 (断点续传)", type="primary", use_container_width=True):
    if not uploaded_files or not api_key:
        st.error("请上传文件并填写 API Key！")
    else:
        # 初始化未记录的文件状态
        for f in uploaded_files:
            if f.name not in st.session_state.doc_states:
                st.session_state.doc_states[f.name] = {
                    'p1_done': False, 'p2_done': False, 'p3_done': False,
                    'paras': [], 'pairs': [], 'auto_terms': {},
                    'p1_doc': None, 'p2_doc': None, 'p3_doc': None, 'p3_md': None, 'term_excel': None
                }
        
        overall_bar = st.progress(0)
        
        for file_idx, uploaded_file in enumerate(uploaded_files):
            filename = uploaded_file.name
            state = st.session_state.doc_states[filename]
            
            # 如果全部完成，直接跳过
            if state['p1_done'] and state['p2_done'] and (not enable_report or state['p3_done']):
                overall_bar.progress((file_idx + 1) / len(uploaded_files))
                continue

            try:
                with st.status(f"⚙️ 正在处理: {filename}", expanded=True) as status:
                    
                    # ---------------- 阶段一：排版清洗 ----------------
                    if not state['p1_done']:
                        status.update(label="【阶段一】AI 智能排版与断句清洗...", state="running")
                        paragraphs = []
                        file_bytes = uploaded_file.read()
                        
                        if filename.lower().endswith(".pdf"):
                            doc_pdf = fitz.open(stream=file_bytes, filetype="pdf")
                            raw_chunks, current_chunk = [], ""
                            for page in doc_pdf:
                                text = page.get_text("text").strip()
                                if text: current_chunk += text + "\n\n"
                                if len(current_chunk) > 2500:
                                    raw_chunks.append(current_chunk)
                                    current_chunk = ""
                            if current_chunk: raw_chunks.append(current_chunk)
                            doc_pdf.close()

                            st.info(f"📄 初步提取完毕，共 {len(raw_chunks)} 个区块，开始智能排版...")
                            sys_p1 = "你是一个学术排版专家。剔除页眉页脚、合并换行截断的句子。严格返回JSON数组（List[str]）。"
                            
                            for idx, chunk in enumerate(raw_chunks):
                                st.caption(f"📡 清洗区块 {idx+1}/{len(raw_chunks)}...")
                                success = False
                                for attempt in range(3):
                                    try:
                                        result_text = call_llm(ai_provider, api_key, sys_p1, f"文本：\n{chunk}")
                                        json_match = re.search(r'\[.*\]', result_text, re.DOTALL)
                                        if json_match:
                                            for p in json.loads(json_match.group(0)):
                                                for sub_p in re.split(r'\n+', clean_xml_chars(str(p))):
                                                    if len(sub_p.strip()) > 5: paragraphs.append(sub_p.strip())
                                            success = True
                                            time.sleep(1)
                                            break
                                    except Exception as e:
                                        if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e): time.sleep(10)
                                        else: break
                                if not success:
                                    for sub_p in re.split(r'\n+', clean_xml_chars(chunk)):
                                        if len(sub_p.strip()) > 5: paragraphs.append(sub_p.strip())
                        elif filename.lower().endswith(".docx"):
                            doc_word = Document(io.BytesIO(file_bytes))
                            for p in doc_word.paragraphs:
                                for sub_p in re.split(r'\n+', clean_xml_chars(p.text)):
                                    if len(sub_p.strip()) > 5: paragraphs.append(sub_p.strip())

                        if not paragraphs: raise ValueError("未提取到有效文本")
                        
                        # 留存阶段一成果
                        state['paras'] = paragraphs
                        state['p1_doc'] = paragraphs_to_word(paragraphs)
                        state['p1_done'] = True
                        st.session_state.doc_states[filename] = state
                        st.info(f"✅ 阶段一完成！提取 {len(paragraphs)} 个纯净段落，进度已自动保存。")
                    else:
                        st.info("⏭️ 阶段一已完成，断点跳过。")

                    # ---------------- 阶段 1.5：智能抽取术语 ----------------
                    final_termbase = user_termbase.copy()
                    if auto_term and not state['auto_terms']:
                        status.update(label="【阶段1.5】正在 AI 智能抽取全文核心术语...", state="running")
                        st.caption("🤖 正在从前言样本中提取行业词汇...")
                        extracted = extract_auto_terms(state['paras'], target_lang, ai_provider, api_key)
                        state['auto_terms'] = extracted
                        if extracted:
                            state['term_excel'] = dict_to_excel(extracted)
                            st.success(f"✅ 成功提取 {len(extracted)} 个专属术语！并已注入翻译引擎。")
                        st.session_state.doc_states[filename] = state
                    
                    if state['auto_terms']:
                        final_termbase.update(state['auto_terms'])

                    # ---------------- 阶段二：双语翻译 ----------------
                    if not state['p2_done']:
                        status.update(label="【阶段二】双语翻译与术语严格注入...", state="running")
                        bilingual_pairs = state['pairs'] # 读取之前的断点进度
                        start_idx = len(bilingual_pairs) # 断点续传的核心！
                        
                        out_doc = Document()
                        table = out_doc.add_table(rows=1, cols=2)
                        table.style = 'Table Grid'
                        
                        # 恢复已翻译的部分到文档内存
                        for pair in bilingual_pairs:
                            row = table.add_row().cells
                            row[0].text = pair['source']
                            row[1].text = pair['target']
                        
                        term_prompt = "\n【强制术语】：\n" + "\n".join([f"- {k} -> {v}" for k, v in final_termbase.items()]) if final_termbase else ""
                        sys_prompt = f"你是一个学术翻译专家，请翻译成{target_lang}。纯作者信息保留原文。{term_prompt}"
                        
                        p2_progress = st.progress(start_idx / len(state['paras']))
                        for i in range(start_idx, len(state['paras'])):
                            para = state['paras'][i]
                            st.caption(f"🌍 正在翻译第 {i+1}/{len(state['paras'])} 段...")
                            trans = ""
                            for attempt in range(3):
                                try:
                                    trans = call_llm(ai_provider, api_key, sys_prompt, para, temperature=0.3)
                                    time.sleep(1)
                                    break
                                except Exception as e:
                                    if '429' in str(e) or 'RESOURCE_EXHAUSTED' in str(e): time.sleep(15)
                                    else: raise e
                            
                            clean_trans = clean_xml_chars(trans).replace('\n', ' ')
                            clean_para = para.replace('\n', ' ')
                            
                            row = table.add_row().cells
                            row[0].text = clean_para
                            row[1].text = clean_trans
                            bilingual_pairs.append({"source": clean_para, "target": clean_trans})
                            p2_progress.progress((i + 1) / len(state['paras']))
                            
                            # 每翻译一段就保存一下状态字典（真正的极细粒度断点）
                            state['pairs'] = bilingual_pairs
                            st.session_state.doc_states[filename] = state

                        doc_io = io.BytesIO()
                        out_doc.save(doc_io)
                        doc_io.seek(0)
                        state['p2_doc'] = doc_io
                        state['p2_done'] = True
                        st.session_state.doc_states[filename] = state
                        st.info("✅ 阶段二完成！翻译对照表已生成。")
                    else:
                        st.info("⏭️ 阶段二已完成，断点跳过。")

                    # ---------------- 阶段三：报告生成 ----------------
                    if enable_report and not state['p3_done']:
                        status.update(label=f"【阶段三】基于《{translation_theory}》生成报告...", state="running")
                        report_md = generate_mti_report(state['pairs'], final_termbase, translation_theory, ai_provider, api_key, status_placeholder=status)
                        
                        if "生成失败" not in report_md:
                            state['p3_md'] = report_md
                            state['p3_doc'] = markdown_to_word(report_md, translation_theory)
                            state['p3_done'] = True
                            st.session_state.doc_states[filename] = state
                            st.info("✅ 阶段三完成！报告已生成。")
                        else:
                            st.error("❌ 报告生成超时，请重新点击“继续处理”重试。")
                            raise ValueError("Report Timeout")
                    elif enable_report:
                        st.info("⏭️ 阶段三已完成，断点跳过。")

                    status.update(label=f"🎉 {filename} 全部流程圆满完成！", state="complete")
            except Exception as e:
                st.error(f"⚠️ {filename} 处理中断: {str(e)}。不要慌，进度已保存，再次点击顶部按钮即可继续！")
            
            overall_bar.progress((file_idx + 1) / len(uploaded_files))

# ================= 动态渲染过程资产面板 (随时下载) =================
if st.session_state.doc_states:
    st.divider()
    st.header("📦 项目过程资产沉淀")
    for filename, state in st.session_state.doc_states.items():
        with st.expander(f"📁 资产面板: {filename}", expanded=True):
            col_d1, col_d2, col_d3, col_d4 = st.columns(4)
            
            with col_d1:
                if state['p1_doc']:
                    st.download_button(f"📥 1. 洗净后原文", state['p1_doc'], file_name=f"阶段1_清洗原文_{filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"d1_{filename}", use_container_width=True)
            with col_d2:
                if state['term_excel']:
                    st.download_button(f"🧠 1.5 提取术语库", state['term_excel'], file_name=f"自动抽词库_{filename}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", key=f"dt_{filename}", use_container_width=True)
            with col_d3:
                if state['p2_doc']:
                    st.download_button(f"📥 2. 双语对照表", state['p2_doc'], file_name=f"阶段2_双语对照_{filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"d2_{filename}", use_container_width=True)
            with col_d4:
                if state['p3_doc']:
                    st.download_button(f"📝 3. 翻译实践报告", state['p3_doc'], file_name=f"阶段3_实践报告_{filename}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"d3_{filename}", use_container_width=True)
            
            if state['p3_md']:
                st.markdown(state['p3_md'])